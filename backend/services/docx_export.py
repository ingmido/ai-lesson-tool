"""
Builds .docx files that mirror the layout of the two supplied templates:
  - Template 1: កិច្ចតែងការបង្រៀន (lesson plan) — header/meta block, objectives,
    materials, then a 3-column activity table split into 5 numbered steps.
  - Template 2: កម្មវិធីសិក្សា (curriculum distribution) — a wide table with
    month / week+hours / content / outcome / page columns.

Slide & test exports use simpler, clean layouts (no existing template given).

NOTE: A Khmer Unicode font (e.g. "Khmer OS Battambang") must be installed on
the machine that opens these files for correct rendering; python-docx just
tags the font name, it doesn't embed it.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KHMER_FONT = "Khmer OS Battambang"


def _set_khmer_font(run, size=11, bold=False):
    run.font.name = KHMER_FONT
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), KHMER_FONT)
    rFonts.set(qn("w:cs"), KHMER_FONT)


def _para(doc, text="", size=11, bold=False, align=None, space_after=4):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_khmer_font(run, size=size, bold=bold)
    return p


def _cell_text(cell, text, size=10, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text) if text is not None else "")
    _set_khmer_font(run, size=size, bold=bold)


def _shade_cell(cell, hex_color="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_lesson_plan_docx(data, output_path, teacher_name="", school_name=""):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    _para(doc, "កិច្ចតែងការបង្រៀន", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    meta_rows = [
        ("កាលបរិច្ឆេទ", data.get("date", "")),
        ("ថ្នាក់ទី", data.get("grade", "")),
        ("មុខវិជ្ជា", data.get("subject", "")),
        ("ជំពូកទី", data.get("chapter", "")),
        ("មេរៀនទី", data.get("lesson", "")),
        ("សិក្សាអំពី", data.get("topic", "")),
        ("រយៈពេល", data.get("duration", "")),
        ("គោលវិធី", data.get("approach", "")),
        ("វិធីសាស្ត្រ", data.get("method", "")),
    ]
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.autofit = True
    for label, value in meta_rows:
        row = meta_table.add_row()
        _cell_text(row.cells[0], label, size=11, bold=True)
        row.cells[0].width = Cm(4)
        _cell_text(row.cells[1], value, size=11)

    doc.add_paragraph()
    _para(doc, "១. វត្ថុបំណង៖", size=12, bold=True)
    obj = data.get("objectives", {}) or {}
    _para(doc, f"វិជ្ជាទស្សន៍៖ {obj.get('knowledge', '')}", size=11)
    _para(doc, f"បំណិន៖ {obj.get('skills', '')}", size=11)
    _para(doc, f"ចរិយា៖ {obj.get('attitude', '')}", size=11)

    doc.add_paragraph()
    _para(doc, "២. សម្ភារៈឧបទេស៖", size=12, bold=True)
    mat = data.get("materials", {}) or {}
    _para(doc, f"ឯកសារ៖ {mat.get('documents', '')}", size=11)
    _para(doc, f"ឧបករណ៍៖ {mat.get('tools', '')}", size=11)

    doc.add_paragraph()
    _para(doc, "៣. សកម្មភាពបង្រៀន និងសិក្សា៖", size=12, bold=True)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    headers = ["សកម្មភាពគ្រូ", "ខ្លឹមសារ", "សកម្មភាពសិស្ស"]
    widths = [Cm(5.3), Cm(6.7), Cm(5.3)]
    for i, h in enumerate(headers):
        _cell_text(hdr[i], h, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(hdr[i])
        hdr[i].width = widths[i]

    for step in data.get("activities", []):
        title_row = table.add_row().cells
        merged = title_row[0].merge(title_row[1]).merge(title_row[2])
        step_label = step.get("step_title", "")
        time_label = step.get("time", "")
        _cell_text(merged, f"{step_label} ({time_label})", size=11, bold=True)
        _shade_cell(merged, "EFEFEF")

        row = table.add_row().cells
        _cell_text(row[0], step.get("teacher_activity", ""), size=10)
        _cell_text(row[1], step.get("content", ""), size=10)
        _cell_text(row[2], step.get("student_activity", ""), size=10)
        for c, w in zip(row, widths):
            c.width = w

    doc.add_paragraph()
    doc.add_paragraph()
    sign = doc.add_table(rows=1, cols=2)
    _cell_text(sign.rows[0].cells[0], "", size=11)
    right = sign.rows[0].cells[1]
    right.text = ""
    p1 = right.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("អ្នករៀបរៀង និងបង្រៀន")
    _set_khmer_font(r1, size=11, bold=True)
    p2 = right.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(teacher_name or "……………………………")
    _set_khmer_font(r2, size=11)

    doc.save(output_path)
    return output_path


def build_curriculum_docx(data, output_path):
    doc = Document()
    section = doc.sections[0]
    section.orientation = 1  # landscape hint; set dims below
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    _para(doc, "កម្មវិធីសិក្សា", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    subtitle = f"{data.get('subject', '')} - {data.get('grade', '')} - {data.get('school_year', '')}"
    _para(doc, subtitle, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["ខែ", "សប្តាហ៍ (ម៉ោង)", "ខ្លឹមសារមេរៀន", "លទ្ធផលការសិក្សា", "រ.ក", "រ.គ"]
    widths = [Cm(2.2), Cm(2.8), Cm(9.5), Cm(9.5), Cm(1.5), Cm(1.5)]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _cell_text(hdr[i], h, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(hdr[i])
        hdr[i].width = widths[i]

    for row_data in data.get("rows", []):
        row = table.add_row().cells
        _cell_text(row[0], row_data.get("month", ""), size=10, bold=True)
        _cell_text(row[1], row_data.get("week_hours", ""), size=10)
        _cell_text(row[2], row_data.get("lesson_content", ""), size=10)
        _cell_text(row[3], row_data.get("learning_outcome", ""), size=10)
        _cell_text(row[4], row_data.get("page_rk", ""), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(row[5], row_data.get("page_rg", ""), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        for c, w in zip(row, widths):
            c.width = w

    doc.save(output_path)
    return output_path


def build_slide_docx(data, output_path):
    """Text outline export of the slide deck (the real visual deck is the .pptx export)."""
    doc = Document()
    _para(doc, data.get("title", "ស្លាយបង្រៀន"), size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, data.get("subtitle", ""), size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    for i, s in enumerate(data.get("slides", []), start=1):
        heading = s.get("heading", "")
        _para(doc, f"ស្លាយទី {i}៖ {heading}", size=13, bold=True, space_after=4)

        stype = s.get("type", "")
        if stype == "bullets":
            for item in s.get("items", []):
                _para(doc, f"•  {item.get('title', '')} — {item.get('desc', '')}", size=11, space_after=2)
        elif stype == "two_column":
            for p in s.get("paragraphs", []):
                _para(doc, p, size=11, space_after=2)
            if s.get("stat_number"):
                _para(doc, f"({s.get('stat_number')} — {s.get('stat_label', '')})", size=10, space_after=2)
        elif stype == "cards":
            for c in s.get("cards", []):
                _para(doc, f"•  {c.get('title', '')} — {c.get('desc', '')}", size=11, space_after=2)
        elif stype == "timeline":
            for st in s.get("stages", []):
                _para(doc, f"•  {st.get('title', '')} ({st.get('subtitle', '')}) — {st.get('desc', '')}",
                      size=11, space_after=2)
        elif stype == "table":
            cols = s.get("columns", [])
            if cols:
                _para(doc, " | ".join(str(c) for c in cols), size=11, bold=True, space_after=2)
            for row in s.get("rows", []):
                _para(doc, " | ".join(str(v) for v in row), size=10, space_after=1)
        elif stype == "compare":
            left, right = s.get("left", {}), s.get("right", {})
            _para(doc, f"{left.get('title', '')}: " + "; ".join(left.get("bullets", [])), size=11, space_after=2)
            _para(doc, f"{right.get('title', '')}: " + "; ".join(right.get("bullets", [])), size=11, space_after=2)
        elif stype == "closing":
            for j, q in enumerate(s.get("questions", []), start=1):
                _para(doc, f"{j}. {q}", size=11, space_after=2)
        elif s.get("subheading"):
            _para(doc, s.get("subheading", ""), size=11, space_after=2)

        doc.add_paragraph()

    doc.save(output_path)
    return output_path


def build_test_docx(data, output_path):
    doc = Document()
    _para(doc, data.get("title", "តេស្តស្តង់ដារ"), size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    if data.get("instructions"):
        _para(doc, data["instructions"], size=11, space_after=2)
    if data.get("duration"):
        _para(doc, f"រយៈពេល៖ {data['duration']}", size=11, space_after=10)

    for section_data in data.get("sections", []):
        _para(doc, section_data.get("section_title", ""), size=13, bold=True, space_after=6)
        for q in section_data.get("questions", []):
            _para(doc, f"{q.get('number', '')}. {q.get('question', '')}", size=11, space_after=2)
            for choice in q.get("choices", []) or []:
                _para(doc, f"     {choice}", size=11, space_after=1)
            doc.add_paragraph()

    doc.add_page_break()
    _para(doc, "ចម្លើយគំរូ", size=14, bold=True, space_after=8)
    for section_data in data.get("sections", []):
        for q in section_data.get("questions", []):
            _para(doc, f"{q.get('number', '')}. {q.get('answer', '')}", size=11, space_after=2)

    doc.save(output_path)
    return output_path


BUILDERS = {
    "lesson_plan": build_lesson_plan_docx,
    "curriculum": build_curriculum_docx,
    "slide": build_slide_docx,
    "test": build_test_docx,
}


def export_to_docx(tool_type, data, output_path, **kwargs):
    builder = BUILDERS.get(tool_type)
    if not builder:
        raise ValueError(f"គ្មាន docx builder សម្រាប់ {tool_type}")
    return builder(data, output_path, **kwargs) if tool_type == "lesson_plan" else builder(data, output_path)


def _js_string(s):
    """Escape a Python string for safe embedding inside a JS single-quoted string literal."""
    if s is None:
        s = ""
    return (
        str(s)
        .replace("\\", "\\\\")
        .replace("'", "\\'")
        .replace("\n", "\\n")
        .replace("\r", "")
    )


def build_google_form_script(data):
    """
    Builds a Google Apps Script (.gs) source that, when pasted into
    script.google.com and run, creates a Google Form quiz matching this
    test's questions. This avoids needing server-side Google OAuth —
    the script runs under the teacher's own Google account.
    """
    title = _js_string(data.get("title", "តេស្ត"))
    instructions = _js_string(data.get("instructions", ""))

    lines = []
    lines.append("function createFormFromTest() {")
    lines.append(f"  var form = FormApp.create('{title}');")
    lines.append("  form.setIsQuiz(true);")
    if instructions:
        lines.append(f"  form.setDescription('{instructions}');")
    lines.append("")

    q_num = 0
    for section in data.get("sections", []):
        section_title = _js_string(section.get("section_title", ""))
        if section_title:
            lines.append(f"  form.addSectionHeaderItem().setTitle('{section_title}');")
        for q in section.get("questions", []):
            q_num += 1
            question_text = _js_string(q.get("question", ""))
            choices = q.get("choices")
            answer = _js_string(q.get("answer", ""))

            if choices:
                lines.append(f"  var q{q_num} = form.addMultipleChoiceItem();")
                lines.append(f"  q{q_num}.setTitle('{question_text}');")
                choice_lines = ", ".join(f"q{q_num}.createChoice('{_js_string(c)}')" for c in choices)
                lines.append(f"  q{q_num}.setChoices([{choice_lines}]);")
                lines.append(f"  q{q_num}.setPoints(1);")
            else:
                lines.append(f"  var q{q_num} = form.addParagraphTextItem();")
                lines.append(f"  q{q_num}.setTitle('{question_text}');")
                lines.append(f"  q{q_num}.setPoints(1);")
            if answer:
                lines.append(f"  // ចម្លើយត្រឹមត្រូវ៖ {answer}")
            lines.append("")

    lines.append("  Logger.log('Form created: ' + form.getEditUrl());")
    lines.append("  Logger.log('Share link: ' + form.getPublishedUrl());")
    lines.append("}")
    return "\n".join(lines)
